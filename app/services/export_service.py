from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models import LearningLog
from app.services import ai_service
from app.services.log_service import normalize_username


def list_logs_for_export(
    db: Session,
    username: str,
    start_date: date,
    end_date: date,
) -> list[LearningLog]:
    return (
        db.query(LearningLog)
        .filter(
            LearningLog.username == normalize_username(username),
            LearningLog.date >= start_date,
            LearningLog.date <= end_date,
        )
        .order_by(LearningLog.date.asc(), LearningLog.id.asc())
        .all()
    )


def extract_note(content: str) -> str:
    value = content or ""
    if value.startswith("心情：") and "\n\n" in value:
        return value.split("\n\n", 1)[1].strip()
    return value.strip()


def build_summary_context(logs: list[LearningLog]) -> str:
    if not logs:
        return "该时间范围内暂无学习记录。"
    chunks = []
    for row in logs:
        chunks.append(
            "\n".join(
                [
                    f"日期：{row.date}",
                    f"今日计划：{row.tasks or '暂无'}",
                    f"待解决事项：{row.problems or '暂无'}",
                    f"今日笔记：{extract_note(row.content) or '暂无'}",
                ]
            )
        )
    return "\n\n---\n\n".join(chunks)[-12000:]


def generate_ai_summary(logs: list[LearningLog], start_date: date, end_date: date) -> str:
    if not logs:
        return "该时间范围内暂无学习记录，无法形成阶段性总结。"
    messages = [
        {
            "role": "system",
            "content": (
                "你是个人学习管理系统的阶段总结助手。"
                "请根据用户的学习日志生成简洁、具体、可执行的阶段总结。"
                "不要总结今日计划整体完成情况。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"时间范围：{start_date} 至 {end_date}\n\n"
                f"学习日志：\n{build_summary_context(logs)}\n\n"
                "请按以下四部分输出：\n"
                "1. 本阶段主要做了什么事情\n"
                "2. 本阶段暴露的问题和待解决事项\n"
                "3. 未来需要改进什么\n"
                "4. 后续应该朝哪个方向规划\n"
            ),
        },
    ]
    try:
        return ai_service.remove_emoji(ai_service.call_chat_model(messages))
    except Exception as exc:
        return f"AI 总结生成失败：{exc}"


def set_document_font(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(11)


def add_text_block(document: Document, title: str, text: str) -> None:
    document.add_heading(title, level=3)
    paragraph = document.add_paragraph()
    paragraph.add_run(text or "暂无")


def build_export_docx(
    db: Session,
    username: str,
    start_date: date,
    end_date: date,
) -> bytes:
    logs = list_logs_for_export(db, username, start_date, end_date)
    summary = generate_ai_summary(logs, start_date, end_date)

    document = Document()
    set_document_font(document)
    document.add_heading("个人学习记录导出", level=1)
    document.add_paragraph(f"用户：{normalize_username(username)}")
    document.add_paragraph(f"时间范围：{start_date} 至 {end_date}")

    document.add_heading("一、阶段 AI 总结", level=2)
    for line in summary.splitlines():
        if line.strip():
            document.add_paragraph(line.strip())

    document.add_heading("二、每日详细记录", level=2)
    if not logs:
        document.add_paragraph("该时间范围内暂无学习记录。")
    for row in logs:
        document.add_heading(str(row.date), level=2)
        add_text_block(document, "今日计划", row.tasks or "暂无")
        add_text_block(document, "待解决事项", row.problems or "暂无")
        add_text_block(document, "今日笔记", extract_note(row.content) or "暂无")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
